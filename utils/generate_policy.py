from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def generate_chinese_description(data):
    def to_money(s):
        return f"${int(float(s)):,}" if s else "未提供"

    def build_uninsured_desc(bi, pd, deductible):
        desc = []
        if bi:
            split = bi.split('/')
            desc.append(f"赔偿你和乘客医疗费 {to_money(split[0])}/人，一场事故最多 {to_money(split[1])}")
        if pd:
            desc.append(f"赔偿自己车辆 {to_money(pd)}（自付额 {to_money(deductible or '250')}）")
        return '\n'.join(desc) if desc else "没有选择该项目"

    fields = {
        "责任险": {
            "selected": "✅" if data["liability"]["selected"] else "❌",
            "desc": f"赔偿对方医疗费 {to_money(data['liability']['bodily_injury'])}/人\n"
                    f"一场事故最多 {to_money(data['liability']['bodily_injury'].split('/')[1])}\n"
                    f"赔偿对方车辆 {to_money(data['liability']['property_damage'])}"
        },
        "无保险驾驶者保障": {
            "selected": "✅" if data["uninsured_motorist"]["selected"] else "❌",
            "desc": build_uninsured_desc(
                data["uninsured_motorist"]["bodily_injury"],
                data["uninsured_motorist"]["property_damage"],
                data["uninsured_motorist"]["property_deductible"]
            )
        },
        "医疗费用": {
            "selected": "✅" if data["medical_payment"]["selected"] else "❌",
            "desc": f"赔偿自己和自己车上乘客在事故中受伤的医疗费每人{to_money(data['medical_payment']['limit'])}"
                    if data["medical_payment"]["selected"] else "没有选择该项目"
        },
        "人身损失": {
            "selected": "✅" if data["personal_injury"]["selected"] else "❌",
            "desc": f"赔偿自己和自己车上乘客在事故中受伤的医疗费、误工费和精神损失费每人{to_money(data['personal_injury']['limit'])}"
                    if data["personal_injury"]["selected"] else "没有选择该项目"
        }
    }
    return fields

def generate_policy_doc(template_path, output_path, data):
    doc = Document(template_path)

    def set_cell_text(cell, text, align="center", font_size=12, bold=False):
        cell.text = text
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == "center" else WD_ALIGN_PARAGRAPH.LEFT
        run = p.runs[0]
        run.font.size = Pt(font_size)
        run.bold = bold

    description = generate_chinese_description(data)
    table = doc.tables[0]

    for i, label in enumerate(["责任险", "无保险驾驶者保障", "医疗费用", "人身损失"]):
        set_cell_text(table.cell(i+1, 1), description[label]["selected"], font_size=16)
        set_cell_text(table.cell(i+1, 2), description[label]["desc"], align="left", font_size=11)

    # 后续车辆保障部分插入逻辑可另外添加

    doc.save(output_path)
